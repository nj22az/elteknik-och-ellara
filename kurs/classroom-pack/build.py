#!/usr/bin/env python3
"""Build A4 PDF + Word Classroom packs for Elteknik och ellära, weeks 1–9.

Sources are the existing elevblad, lektioner and classroom posts.
Does not invent legal text. Does not invent elevblad when none exist.
Print: IBM Plex Sans, ink on paper, 0 radius.
"""
from __future__ import annotations

import html
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips, Emu
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

ROOT = Path(__file__).resolve().parents[2]
FONTS = ROOT / "pptx" / "fonts"
OUT = ROOT / "kurs" / "classroom-pack"
INK = "111111"
PAPER = "FFFFFF"
INK_RGB = RGBColor(0x11, 0x11, 0x11)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)

FONT_REG = FONTS / "IBMPlexSans-Regular.ttf"
FONT_MED = FONTS / "IBMPlexSans-Medium.ttf"
FONT_SEM = FONTS / "IBMPlexSans-SemiBold.ttf"


def font_uri(p: Path) -> str:
    return p.resolve().as_uri()


CSS_TEXT = f"""
@font-face {{
  font-family: "IBM Plex Sans";
  src: url("{font_uri(FONT_REG)}");
  font-weight: 400;
  font-style: normal;
}}
@font-face {{
  font-family: "IBM Plex Sans";
  src: url("{font_uri(FONT_MED)}");
  font-weight: 500;
  font-style: normal;
}}
@font-face {{
  font-family: "IBM Plex Sans";
  src: url("{font_uri(FONT_SEM)}");
  font-weight: 600;
  font-style: normal;
}}
@font-face {{
  font-family: "IBM Plex Sans";
  src: url("{font_uri(FONT_SEM)}");
  font-weight: 700;
  font-style: normal;
}}
:root {{
  --ink: #{INK};
  --paper: #{PAPER};
}}
* {{ box-sizing: border-box; border-radius: 0 !important; }}
html, body {{
  margin: 0;
  padding: 0;
  background: var(--paper);
  color: var(--ink);
  font-family: "IBM Plex Sans", "Liberation Sans", "DejaVu Sans", sans-serif;
  font-size: 10.5pt;
  line-height: 1.38;
  font-variant-numeric: tabular-nums lining-nums;
}}
@page {{
  size: A4;
  margin: 12mm 14mm 16mm 14mm;
  @bottom-left {{
    content: element(foot-left);
    color: #{INK};
    font-family: "IBM Plex Sans", sans-serif;
    font-size: 7.5pt;
    font-weight: 600;
    letter-spacing: 0.04em;
  }}
  @bottom-right {{
    content: counter(page) " / " counter(pages);
    color: #{INK};
    font-family: "IBM Plex Sans", sans-serif;
    font-size: 7.5pt;
    font-variant-numeric: tabular-nums;
  }}
}}
.foot-left {{
  position: running(foot-left);
  font-weight: 600;
  font-size: 7.5pt;
  letter-spacing: 0.04em;
}}
.teacher-banner {{
  background: #{INK};
  color: #{PAPER};
  font-weight: 600;
  font-size: 8.5pt;
  letter-spacing: 0.08em;
  text-transform: uppercase;
  padding: 5px 8px;
  margin: 0 0 10px 0;
}}
.chrome {{
  display: flex;
  align-items: center;
  gap: 8px;
  margin: 0 0 6px 0;
  flex-wrap: wrap;
}}
.loz, .vsq {{
  display: inline-block;
  background: #{INK};
  color: #{PAPER};
  font-weight: 600;
  letter-spacing: 0.08em;
  font-size: 9pt;
  padding: 2px 7px;
  line-height: 1.2;
}}
.vsq {{ letter-spacing: 0.06em; }}
.chrome-name {{
  font-weight: 600;
  font-size: 10.5pt;
  letter-spacing: 0.01em;
}}
.chrome-meta {{
  margin-left: auto;
  font-size: 8.5pt;
  font-weight: 500;
}}
h1 {{
  font-size: 18pt;
  font-weight: 600;
  margin: 4px 0 2px 0;
  line-height: 1.15;
  page-break-after: avoid;
}}
.kicker {{
  font-size: 9pt;
  font-weight: 500;
  margin: 0 0 8px 0;
}}
h2 {{
  font-size: 12pt;
  font-weight: 600;
  margin: 9px 0 4px 0;
  padding-top: 3px;
  border-top: 1px solid #{INK};
  page-break-after: avoid;
}}
h2.first {{ border-top: none; padding-top: 0; margin-top: 10px; }}
h3 {{
  font-size: 10.5pt;
  font-weight: 600;
  margin: 10px 0 4px 0;
  page-break-after: avoid;
}}
p {{ margin: 0 0 6px 0; }}
strong {{ font-weight: 600; }}
em {{ font-style: italic; }}
code {{
  font-family: "IBM Plex Sans", sans-serif;
  font-weight: 600;
  font-size: 0.95em;
  letter-spacing: 0.02em;
}}
ul, ol {{ margin: 0 0 8px 0; padding-left: 1.2em; }}
li {{ margin: 0 0 3px 0; }}
.meta {{
  display: flex;
  gap: 18px;
  flex-wrap: wrap;
  margin: 8px 0 10px 0;
  font-size: 10.5pt;
}}
.meta .field {{ flex: 1; min-width: 180px; }}
.line {{
  border-bottom: 1px solid #{INK};
  display: inline-block;
  min-width: 220px;
  min-height: 1.15em;
}}
.line.short {{ min-width: 90px; }}
.line.full {{ display: block; min-width: 100%; height: 1.5em; margin: 4px 0 8px 0; }}
.checks {{ margin: 4px 0 8px 0; }}
.check {{
  display: inline-block;
  margin-right: 18px;
  font-weight: 500;
}}
.box {{
  border: 1px solid #{INK};
  margin: 0 0 8px 0;
  page-break-inside: avoid;
}}
.box.heavy {{ border-width: 2px; }}
.box .tab {{
  display: inline-block;
  background: #{INK};
  color: #{PAPER};
  font-weight: 600;
  font-size: 8.5pt;
  letter-spacing: 0.08em;
  padding: 2px 8px;
}}
.box .body {{ padding: 7px 10px 8px 10px; }}
.fara .head {{
  background: #{INK};
  color: #{PAPER};
  font-weight: 600;
  letter-spacing: 0.08em;
  font-size: 8.5pt;
  padding: 4px 10px;
}}
.fara .body {{ padding: 7px 10px 8px 10px; }}
.row2 {{
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 12px;
  margin: 8px 0 10px 0;
}}
.row3 {{
  display: grid;
  grid-template-columns: 1fr 1fr 1fr;
  gap: 8px;
  margin: 8px 0 10px 0;
}}
.choice {{
  border: 2px solid #{INK};
  text-align: center;
  padding: 8px 8px;
  page-break-inside: avoid;
}}
.choice .lab {{
  font-weight: 600;
  font-size: 16pt;
  letter-spacing: 0.12em;
  margin-bottom: 6px;
}}
.choice.filled {{
  background: #{INK};
  color: #{PAPER};
}}
.card {{
  border: 1px solid #{INK};
  text-align: center;
  padding: 10px 6px;
  font-weight: 600;
  font-size: 11pt;
  page-break-inside: avoid;
}}
.card.circled {{
  outline: 2px solid #{INK};
  outline-offset: -5px;
}}
.diag {{
  border: 1px solid #{INK};
  padding: 8px 10px 8px 10px;
  margin: 0 0 6px 0;
  page-break-inside: avoid;
  text-align: center;
}}
.housing {{
  border: 1px solid #{INK};
  width: 70%;
  margin: 0 auto 0 auto;
  padding: 10px 8px 12px 8px;
}}
.housing .hl {{
  font-weight: 600;
  letter-spacing: 0.08em;
  font-size: 10pt;
  margin-bottom: 8px;
}}
.valve {{
  border: 1px solid #{INK};
  width: 42%;
  margin: 0 auto;
  padding: 8px 4px;
  font-weight: 600;
  font-size: 9pt;
  letter-spacing: 0.04em;
}}
.stem {{
  width: 1px;
  height: 14px;
  background: #{INK};
  margin: 0 auto;
}}
.water {{
  font-weight: 500;
  letter-spacing: 0.12em;
  font-size: 9pt;
}}
.hand {{
  border: 1px solid #{INK};
  display: inline-block;
  padding: 5px 18px;
  font-weight: 600;
  letter-spacing: 0.08em;
  font-size: 10pt;
}}
.hull {{
  border-top: 3px solid #{INK};
  margin: 0 8%;
  padding-top: 4px;
  font-weight: 600;
  letter-spacing: 0.14em;
  font-size: 10pt;
}}
.disk {{
  display: inline-block;
  width: 16px;
  height: 16px;
  background: #{INK};
  color: #{PAPER};
  font-weight: 600;
  font-size: 9pt;
  line-height: 16px;
  text-align: center;
  border-radius: 50% !important;
  margin-right: 4px;
}}
.ann {{
  font-size: 8.5pt;
  font-weight: 600;
  letter-spacing: 0.02em;
}}
.fillrow {{
  margin: 8px 0 4px 0;
  font-weight: 500;
}}
.fillrow .slot {{
  display: inline-block;
  border-bottom: 1px solid #{INK};
  min-width: 88px;
  text-align: center;
  font-weight: 600;
}}
table.data {{
  width: 100%;
  border-collapse: collapse;
  margin: 0 0 10px 0;
  page-break-inside: avoid;
}}
table.data th, table.data td {{
  border: 1px solid #{INK};
  padding: 4px 7px;
  text-align: left;
  vertical-align: top;
  font-size: 9.5pt;
}}
table.data th {{
  font-weight: 600;
  background: #{PAPER};
}}
table.form {{
  width: 100%;
  border-collapse: collapse;
  margin: 0 0 10px 0;
  border: 2px solid #{INK};
}}
table.form th.form-head {{
  background: #{INK};
  color: #{PAPER};
  font-weight: 600;
  letter-spacing: 0.08em;
  font-size: 8.5pt;
  text-align: left;
  padding: 5px 8px;
  border: none;
}}
table.form td {{
  border-top: 1px solid #{INK};
  padding: 4px 8px;
  vertical-align: top;
}}
table.form td.lab {{
  font-weight: 600;
  width: 26%;
  letter-spacing: 0.01em;
}}
.chain {{
  display: grid;
  grid-template-columns: 28px 1fr;
  gap: 4px 8px;
  align-items: center;
}}
.sq {{
  display: inline-block;
  width: 22px;
  height: 22px;
  border: 1px solid #{INK};
  text-align: center;
  font-weight: 600;
  line-height: 20px;
  font-size: 10pt;
}}
.fig {{
  margin: 8px 0 12px 0;
  page-break-inside: avoid;
}}
.fig img {{
  width: 100%;
  height: auto;
  border: 1px solid #{INK};
}}
.fig .cap {{
  font-size: 8.5pt;
  font-weight: 600;
  margin-top: 4px;
}}
hr.rule {{
  border: none;
  border-top: 1px solid #{INK};
  margin: 8px 0;
}}
.post {{
  border: 1px solid #{INK};
  padding: 10px 12px;
  font-size: 11pt;
  line-height: 1.45;
}}
pre {{
  font-family: "IBM Plex Sans", "Liberation Sans", "DejaVu Sans", sans-serif;
  font-size: 8.5pt;
  line-height: 1.28;
  white-space: pre;
  margin: 0 0 8px 0;
  padding: 6px 8px;
  border: 1px solid #{INK};
}}
a {{ color: #{INK}; text-decoration: none; }}
"""


def inline_html(text: str) -> str:
    if text is None:
        return ""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    out = []
    for p in parts:
        if p.startswith("**") and p.endswith("**") and len(p) >= 4:
            out.append("<strong>" + html.escape(p[2:-2]) + "</strong>")
        elif p.startswith("*") and p.endswith("*") and len(p) >= 2:
            out.append("<em>" + html.escape(p[1:-1]) + "</em>")
        elif p.startswith("`") and p.endswith("`") and len(p) >= 2:
            out.append("<code>" + html.escape(p[1:-1]) + "</code>")
        else:
            out.append(html.escape(p))
    return "".join(out)


def wrap_html(body: str, footer: str, teacher: bool = False) -> str:
    banner = (
        '<div class="teacher-banner">Till lärare · Inte till elev · Inte på elev-GitHub Pages</div>'
        if teacher
        else ""
    )
    return f"""<!DOCTYPE html>
<html lang="sv">
<head>
<meta charset="utf-8"/>
<title>{html.escape(footer)}</title>
</head>
<body>
<div class="foot-left">{html.escape(footer)}</div>
{banner}
{body}
</body>
</html>"""


def chrome(m_code: str, m_name: str, v_code: str, v_name: str, course: str) -> str:
    return f"""
<div class="chrome">
  <span class="loz">{html.escape(m_code)}</span>
  <span class="chrome-name">{html.escape(m_name)}</span>
  <span class="vsq">{html.escape(v_code)}</span>
  <span class="chrome-name">{html.escape(v_name)}</span>
  <span class="chrome-meta">{html.escape(course)}</span>
</div>"""


def callout(kind: str, title: str, body: str) -> str:
    if kind == "fara":
        return f"""
<div class="box fara heavy">
  <div class="head">{html.escape(title)}</div>
  <div class="body">{inline_html(body)}</div>
</div>"""
    return f"""
<div class="box">
  <div class="tab">{html.escape(title)}</div>
  <div class="body">{inline_html(body)}</div>
</div>"""


# ---------------------------------------------------------------------------
# Documents
# ---------------------------------------------------------------------------

def doc_elev_stotvag() -> tuple[str, str, bool]:
    body = chrome("M01", "Elsäkerhet", "V1", "Vecka 1", "Elteknik och ellära 45 p")
    body += "<h1>Stötväg — elevblad</h1>"
    body += '<p class="kicker">På papper. Du går inte till luckan. Du megger inte.</p>'
    body += """
<div class="meta">
  <div class="field">Namn: <span class="line"></span></div>
  <div class="field">Datum: <span class="line short"></span></div>
</div>
<div class="checks">Spår:
  <span class="check">☐ elektriker (land)</span>
  <span class="check">☐ ingenjör (fartyg)</span>
  <span class="check">☐ båda</span>
</div>"""
    body += callout("biljett", "BILJETT", "Det läcker vid pumpen, känn efter var.")
    body += "<p>Pannrum. Pump och magnetventil. Stålskrov.</p>"
    body += callout("obs", "OBS", "Peka stötväg innan du tar i något.")
    body += '<h2 class="first">1. STOP eller GO</h2>'
    body += "<p>Lappen säger: <em>känn efter var.</em> Kryssa <strong>en</strong>. Sen ritar du.</p>"
    body += """
<div class="row2">
  <div class="choice"><div class="lab">STOP</div><div>☐ STOP</div></div>
  <div class="choice"><div class="lab">GO</div><div>☐ GO</div></div>
</div>"""
    body += "<h2>2. Figur — rita stötvägen</h2>"
    body += "<p>Figuren är blank. Rita stötvägen: <strong>hölje → kropp → skrov</strong>.</p>"
    body += """
<div class="diag">
  <div class="housing">
    <div class="hl">PUMPHÖLJE</div>
    <div class="valve">MAGNET-<br>VENTIL</div>
  </div>
  <div class="stem"></div>
  <div class="water">~~~~ vatten ~~~~</div>
  <div class="stem"></div>
  <div class="hand">HAND</div>
  <div class="stem"></div>
  <div class="hull">SKROV</div>
</div>
<p class="fillrow">Stötväg (fyll):
  <span class="slot">&nbsp;</span> →
  <span class="slot">&nbsp;</span> →
  <span class="slot">&nbsp;</span>
</p>"""
    body += "<h2>3. Ring två fel</h2>"
    body += "<p>Tre lappar. Ring <strong>två</strong> fel.</p>"
    body += """
<div class="row3">
  <div class="card">packning</div>
  <div class="card">PE</div>
  <div class="card">utanför IR</div>
</div>"""
    body += "<h2>4. Två spår — en rad</h2>"
    body += "<p>Elektriker: leta inte JFB.</p>"
    body += "<p>Ingenjör: läcker det vid spänningssatt don är det elarbete.</p>"
    return body, "Elteknik och ellära 45 p · V1 · elevblad stötväg", False


def doc_facit_stotvag() -> tuple[str, str, bool]:
    body = chrome("M01", "Elsäkerhet", "V1", "Vecka 1", "Elteknik och ellära 45 p")
    body += "<h1>Stötväg — facit</h1>"
    body += "<p><strong>G:</strong> STOP, stötväg hölje → kropp → skrov, två av tre fel.</p>"
    body += "<p><strong>IG:</strong> GO, handen först, JFB som skydd.</p>"
    body += '<h2 class="first">1. STOP markerat</h2>'
    body += callout("fara", "FARA", "Känn inte efter på spänningssatt don.")
    body += """
<div class="row2">
  <div class="choice filled"><div class="lab">STOP</div><div>☑ STOP</div></div>
  <div class="choice"><div class="lab">GO</div><div>☐ GO</div></div>
</div>"""
    body += "<p><strong>Varför (en rad):</strong> våt magnetventil kan bära spänning. Stötväg hölje → kropp → skrov.</p>"
    body += "<p>GO = IG.</p>"
    body += "<h2>2. Figur — stötväg ifylld</h2>"
    body += """
<div class="diag">
  <div class="housing">
    <div class="hl"><span class="disk">1</span> PUMPHÖLJE <span class="ann">← IN: hölje</span></div>
    <div class="valve">MAGNET-<br>VENTIL</div>
  </div>
  <div class="stem"></div>
  <div class="water">~~~~ vatten ~~~~</div>
  <div class="stem"></div>
  <div><span class="disk">2</span><span class="hand">HAND</span> <span class="ann">← KROPP</span></div>
  <div class="stem"></div>
  <div class="hull">SKROV <span class="ann"><span class="disk">3</span> UT</span></div>
</div>
<p>Stötväg: <strong>hölje → kropp → skrov</strong></p>
<ol>
  <li>In i höljet.</li>
  <li>Genom kroppen.</li>
  <li>Ut i skrovet.</li>
</ol>"""
    body += "<h2>3. Två fel ringade</h2>"
    body += "<p><strong>G = två av tre.</strong> Alla tre duger. Visa minst dessa två:</p>"
    body += """
<div class="row3">
  <div class="card circled">packning</div>
  <div class="card circled">PE</div>
  <div class="card">utanför IR</div>
</div>"""
    body += "<p>Tredje (utanför IR) räknas också. IG: inget ringat, eller JFB som “fel”.</p>"
    body += "<h2>4. Två spår — en rad</h2>"
    body += "<p>Elektriker = leta inte JFB.</p>"
    body += "<p>Ingenjör = läcker det vid spänningssatt don är det elarbete.</p>"
    body += "<p>Samma rad på båda spåren. Inte två prov.</p>"
    body += "<p>Inte skolessä. Inte 1 MΩ. Inte megger.</p>"
    return body, "Elteknik och ellära 45 p · V1 · LÄRARE facit stötväg", True


def megger_form(filled: bool) -> str:
    if filled:
        chain = [
            ("1", "från"),
            ("2", "lås / skylt"),
            ("3", "tvåpol död"),
            ("4", "megger"),
            ("5", "protokoll"),
        ]
        ref_skrov, ref_pe = "☑", "☐"
        a_go, a_stop = "☑", "☐"
        a_till = "GO — planera tillslag EFTER kåpa på"
        b_go, b_stop = "☐", "☑"
        b_till = "STOP — inte spänningssätt, avvikelse"
        ohm_skall, ohm_upp = "☐", "☑"
        chain_intro = "Kedja i ordning:"
    else:
        chain = [
            ("", "megger"),
            ("", "protokoll"),
            ("", "tvåpol död"),
            ("", "från"),
            ("", "lås / skylt"),
        ]
        ref_skrov, ref_pe = "☐", "☐"
        a_go, a_stop = "☐", "☐"
        a_till = ""
        b_go, b_stop = "☐", "☐"
        b_till = ""
        ohm_skall, ohm_upp = "☐", "☐"
        chain_intro = "Kedja — skriv 1–5 i rutorna (rätt ordning):"

    chain_html = '<div class="chain">'
    for n, lab in chain:
        chain_html += f'<div class="sq">{html.escape(n)}</div><div>{html.escape(lab)}</div>'
    chain_html += "</div>"

    till_a = html.escape(a_till) if a_till else "&nbsp;"
    till_b = html.escape(b_till) if b_till else "&nbsp;"

    return f"""
<table class="form">
  <tr><th class="form-head" colspan="2">MEGGERKORT</th></tr>
  <tr><td class="lab">Grupp</td><td>230 V inredning, kabel bytt</td></tr>
  <tr><td class="lab">Tid</td><td>före arbete, avställd grupp</td></tr>
  <tr><td class="lab">{html.escape(chain_intro)}</td><td>{chain_html}</td></tr>
  <tr><td class="lab">Referens</td><td>
    {ref_skrov} skrov / stomme &nbsp;&nbsp; {ref_pe} PE-skena<br>
    <em>Samma grupp hela kortet. Referens är skrov/stomme, inte PE-skena.</em>
  </td></tr>
  <tr><td class="lab">Avläsning A<br>2,4 MΩ</td><td>
    Beslut: {a_go} GO &nbsp;&nbsp; {a_stop} STOP<br>
    Tillslag: <span class="line" style="min-width:70%">{till_a}</span>
  </td></tr>
  <tr><td class="lab">Avläsning B<br>0,4 MΩ</td><td>
    Beslut: {b_go} GO &nbsp;&nbsp; {b_stop} STOP<br>
    Tillslag: <span class="line" style="min-width:70%">{till_b}</span>
  </td></tr>
  <tr><td class="lab">1 MΩ är</td><td>
    {ohm_skall} skall i 5 § &nbsp;&nbsp; {ohm_upp} upplysning
  </td></tr>
</table>"""


def doc_elev_megger() -> tuple[str, str, bool]:
    body = chrome("M02", "Isolering", "V2", "Vecka 2", "Elteknik och ellära 45 p")
    body += "<h1>Meggerkort — elevblad</h1>"
    body += '<p class="kicker">Isolera på papper. Megga inte på spänning.</p>'
    body += """
<div class="meta">
  <div class="field">Namn: <span class="line"></span></div>
  <div class="field">Datum: <span class="line short"></span></div>
</div>
<div class="checks">Spår:
  <span class="check">☐ elektriker (land)</span>
  <span class="check">☐ ingenjör (fartyg)</span>
  <span class="check">☐ båda</span>
</div>"""
    body += callout(
        "biljett",
        "BILJETT",
        "230 V-grupp i inredningen. Kabel bytt. Isolera på papper. Megga inte på spänning.",
    )
    body += callout("fara", "FARA", "Ingen megger på spänningssatt grupp.")
    body += '<h2 class="first">Meggerkort (blankt)</h2>'
    body += megger_form(False)
    body += "<h2>Två spår — en rad</h2>"
    body += "<p>Elektriker: referens är inte PE-skenan hemma.</p>"
    body += "<p>Ingenjör: bara för att maskinen inte låter betyder det inte att den är avställd.</p>"
    body += "<p>Skriv raden för ditt spår:</p>"
    body += '<div class="line full"></div>'
    body += "<h2>VG</h2>"
    body += "<p>Saknas megger på varvet →</p>"
    body += '<div class="line full"></div>'
    body += "<p>(tvåpol + …)</p>"
    return body, "Elteknik och ellära 45 p · V2 · elevblad meggerkort", False


def doc_facit_megger() -> tuple[str, str, bool]:
    body = chrome("M02", "Isolering", "V2", "Vecka 2", "Elteknik och ellära 45 p")
    body += "<h1>Meggerkort — facit</h1>"
    body += "<p><strong>G:</strong> kedja i ordning, referens skrov, 2,4 = GO efter kåpa, 0,4 = STOP + avvikelse, 1 MΩ = upplysning.</p>"
    body += "<p><strong>VG:</strong> saknas megger på varvet → tvåpol + gapflagga.</p>"
    body += "<p><strong>IG:</strong> megger på spänning, PE som referens, 1 MΩ som lag, hoppar lås, 0,4 tillslag.</p>"
    body += '<h2 class="first">Arbetsorder</h2>'
    body += callout(
        "biljett",
        "BILJETT",
        "230 V-grupp i inredningen. Kabel bytt. Isolera på papper. Megga inte på spänning.",
    )
    body += callout("fara", "FARA", "Ingen megger på spänningssatt grupp.")
    body += "<h2>Meggerkort — ifyllt</h2>"
    body += "<p>Kedja <strong>i ordning</strong>. Referens <strong>skrov/stomme</strong>, inte PE.</p>"
    body += megger_form(True)
    body += "<p><strong>2,4 MΩ</strong> = GO att <em>planera</em> tillslag <strong>efter kåpa på</strong>. Inte tillslag med öppen kåpa.</p>"
    body += "<p><strong>0,4 MΩ</strong> = STOP. Ingen tillslag. Skriv avvikelse.</p>"
    body += "<p><strong>1 MΩ</strong> = upplysning, inte lag. Inte “olagligt enligt 5 §”.</p>"
    body += "<h2>Två spår — en rad</h2>"
    body += "<p>Elektriker = referens är inte PE-skenan hemma.</p>"
    body += "<p>Ingenjör = bara för att maskinen inte låter betyder det inte att den är avställd.</p>"
    body += "<h2>VG</h2>"
    body += "<p>Saknas megger på varvet → <strong>tvåpol + gapflagga</strong>. Skriv att meggern saknas. Hitta inte på ett tal. Skippa inte isolation.</p>"
    body += "<p>Inte skolessä. Inte 1 MΩ som skall.</p>"
    return body, "Elteknik och ellära 45 p · V2 · LÄRARE facit meggerkort", True


def parse_md(md: str) -> list[dict]:
    lines = md.replace("\r\n", "\n").split("\n")
    blocks: list[dict] = []
    i = 0
    n = len(lines)
    first_h2 = True

    def take_fence(start: int) -> tuple[int, str]:
        j = start + 1
        buf = []
        while j < n and not lines[j].startswith("```"):
            buf.append(lines[j])
            j += 1
        return j + 1, "\n".join(buf)

    while i < n:
        raw = lines[i]
        s = raw.rstrip()
        if s.strip() == "":
            i += 1
            continue
        if s.startswith("```"):
            i, fence = take_fence(i)
            fence_s = fence.strip()
            kind = None
            title = ""
            body_lines = []
            for ln in fence.splitlines():
                t = ln.strip()
                if "FARA" in t and (t.startswith("█") or t.startswith("┌")):
                    kind = "fara"
                    title = "FARA"
                    continue
                m = re.match(r"^[┌─]+[─\s]*([A-ZÅÄÖ0-9 /]{2,20})", t)
                if t.startswith("┌") and m:
                    kind = "box"
                    title = m.group(1).strip(" ─")
                    continue
                inner = re.sub(r"^[│┃]", "", t)
                inner = re.sub(r"[│┃]$", "", inner)
                inner = inner.strip(" ─")
                inner = inner.strip()
                if inner.startswith("└") or set(inner) <= set("─│└┘┌┐█ "):
                    continue
                if inner:
                    body_lines.append(inner)
            if kind:
                blocks.append({"type": "callout", "kind": kind, "title": title or "OBS", "body": " ".join(body_lines)})
            elif fence_s:
                blocks.append({"type": "pre", "text": fence})
            continue
        if s.strip() in ("---", "***"):
            blocks.append({"type": "hr"})
            i += 1
            continue
        if s.startswith("# "):
            blocks.append({"type": "h1", "text": s[2:].strip()})
            i += 1
            continue
        if s.startswith("## "):
            blocks.append({"type": "h2", "text": s[3:].strip(), "first": first_h2})
            first_h2 = False
            i += 1
            continue
        if s.startswith("### "):
            blocks.append({"type": "h3", "text": s[4:].strip()})
            i += 1
            continue
        if s.startswith("|") and "|" in s[1:]:
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append({"type": "table", "header": rows[0], "rows": rows[1:]})
            continue
        if re.match(r"^[-*] ", s):
            items = []
            while i < n and re.match(r"^[-*] ", lines[i]):
                items.append(re.sub(r"^[-*] ", "", lines[i].rstrip()))
                i += 1
            blocks.append({"type": "ul", "items": items})
            continue
        if re.match(r"^\d+\.\s", s):
            items = []
            while i < n and re.match(r"^\d+\.\s", lines[i]):
                items.append(re.sub(r"^\d+\.\s", "", lines[i].rstrip()))
                i += 1
            blocks.append({"type": "ol", "items": items})
            continue
        para = [s]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if nxt.strip() == "":
                break
            if nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("```") or nxt.strip() in ("---", "***"):
                break
            if re.match(r"^[-*] ", nxt) or re.match(r"^\d+\.\s", nxt):
                break
            if nxt.startswith("**"):
                break
            para.append(nxt)
            i += 1
        blocks.append({"type": "p", "text": " ".join(para)})
    return blocks


def md_body_html(md: str, extra_top: str = "") -> str:
    blocks = parse_md(md)
    parts = [extra_top]
    for b in blocks:
        t = b["type"]
        if t == "h1":
            parts.append(f"<h1>{inline_html(b['text'])}</h1>")
        elif t == "h2":
            cls = ' class="first"' if b.get("first") else ""
            parts.append(f"<h2{cls}>{inline_html(b['text'])}</h2>")
        elif t == "h3":
            parts.append(f"<h3>{inline_html(b['text'])}</h3>")
        elif t == "p":
            parts.append(f"<p>{inline_html(b['text'])}</p>")
        elif t == "hr":
            parts.append('<hr class="rule"/>')
        elif t == "ul":
            lis = "".join(f"<li>{inline_html(x)}</li>" for x in b["items"])
            parts.append(f"<ul>{lis}</ul>")
        elif t == "ol":
            lis = "".join(f"<li>{inline_html(x)}</li>" for x in b["items"])
            parts.append(f"<ol>{lis}</ol>")
        elif t == "table":
            th = "".join(f"<th>{inline_html(c)}</th>" for c in b["header"])
            trs = []
            for row in b["rows"]:
                tds = "".join(f"<td>{inline_html(c)}</td>" for c in row)
                trs.append(f"<tr>{tds}</tr>")
            parts.append(f'<table class="data"><tr>{th}</tr>{"".join(trs)}</table>')
        elif t == "callout":
            parts.append(callout(b["kind"], b["title"], b["body"]))
        elif t == "pre":
            parts.append(f"<pre>{html.escape(b['text'])}</pre>")
    return "\n".join(parts)


def doc_lektion_11() -> tuple[str, str, bool]:
    return doc_lektion(
        ROOT / "kurs/lektioner/1.1-stotar.md",
        ROOT / "bok/figur-1-1-stotvag-ventil.png",
        "M01",
        "Elsäkerhet",
        "V1",
        "Vecka 1",
        "Elteknik och ellära 45 p · V1 · lektion 1.1 stötväg",
        "Figur 1.1 Stötväg — magnetventil (bok). Inte megger. Inte 1 MΩ som skall.",
    )


def doc_lektion_21() -> tuple[str, str, bool]:
    return doc_lektion(
        ROOT / "kurs/lektioner/2.1-isolering.md",
        ROOT / "bok/figur-2-1-isolering-kedja.png",
        "M02",
        "Isolering",
        "V2",
        "Vecka 2",
        "Elteknik och ellära 45 p · V2 · lektion 2.1 isolering",
        "Figur 2.1 Isolering före arbete (bok). Kedja från–lås–prova–megger. 1 MΩ = upplysning.",
    )


def doc_inlagg(week: int) -> tuple[str, str, bool]:
    if week == 1:
        return doc_inlagg_from(
            ROOT / "kurs/lararhandledning/classroom-v1-vecka1.md",
            "M01",
            "Elsäkerhet",
            "V1",
            "Vecka 1",
            "Elteknik och ellära 45 p · V1 · LÄRARE inlägg",
        )
    return doc_inlagg_from(
        ROOT / "kurs/lararhandledning/classroom-v2-vecka2.md",
        "M02",
        "Isolering",
        "V2",
        "Vecka 2",
        "Elteknik och ellära 45 p · V2 · LÄRARE inlägg",
    )


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def write_pdf(html_body: str, footer: str, teacher: bool, dest: Path) -> None:
    font_config = FontConfiguration()
    document = wrap_html(html_body, footer, teacher)
    HTML(string=document, base_url=str(ROOT)).write_pdf(
        dest,
        stylesheets=[CSS(string=CSS_TEXT, font_config=font_config)],
        font_config=font_config,
    )


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def set_run_font(run, size_pt: float, bold: bool = False, color: RGBColor | None = None, white: bool = False):
    run.font.name = "IBM Plex Sans"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = WHITE_RGB if white else (color or INK_RGB)
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "IBM Plex Sans")


def shade_cell(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_borders(cell, sz: str = "8", color: str = INK):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_borders(table, sz: str = "8"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), INK)
        borders.append(el)
    tblPr.append(borders)


def no_table_autofit_indent(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    tblPr.append(jc)


def cell_para(cell, text: str, size: float = 10.5, bold: bool = False, center: bool = False, white: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, white=white)
    return p


def add_inline_runs(p, text: str, size: float = 10.5):
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            r = p.add_run(part[2:-2])
            set_run_font(r, size, bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            r = p.add_run(part[1:-1])
            set_run_font(r, size)
            r.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            r = p.add_run(part[1:-1])
            set_run_font(r, size, bold=True)
        else:
            r = p.add_run(part)
            set_run_font(r, size)


def add_p(doc: Document, text: str, size: float = 10.5, space_after: float = 6, bold: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold:
        r = p.add_run(text)
        set_run_font(r, size, bold=True)
    else:
        add_inline_runs(p, text, size)
    return p


def add_h(doc: Document, text: str, size: float, border_top: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if border_top else 4)
    p.paragraph_format.space_after = Pt(4)
    if border_top:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "6")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), INK)
        pBdr.append(top)
        pPr.append(pBdr)
    r = p.add_run(text)
    set_run_font(r, size, bold=True)
    return p


def add_callout_table(doc: Document, title: str, body: str, fara: bool = False):
    table = doc.add_table(rows=2, cols=1)
    table.autofit = True
    no_table_autofit_indent(table)
    set_table_borders(table, "12" if fara else "8")
    h = table.rows[0].cells[0]
    b = table.rows[1].cells[0]
    shade_cell(h, INK)
    cell_para(h, title, size=9, bold=True, white=True)
    b.text = ""
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, body, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_choice_row(doc: Document, left: tuple[str, str, bool], right: tuple[str, str, bool]):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    set_table_borders(table, "16")
    for i, (lab, sub, filled) in enumerate((left, right)):
        c = table.rows[0].cells[i]
        if filled:
            shade_cell(c, INK)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(lab)
        set_run_font(r, 16, bold=True, white=filled)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(sub)
        set_run_font(r2, 11, bold=True, white=filled)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_three_cards(doc: Document, labels: list[str], circled: list[bool]):
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table, "8")
    for i, lab in enumerate(labels):
        c = table.rows[0].cells[i]
        cell_para(c, ("( " + lab + " )" if circled[i] else lab), size=12, bold=True, center=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_diag(doc: Document, filled: bool):
    table = doc.add_table(rows=5, cols=1)
    set_table_borders(table, "8")
    cells = table.rows
    if filled:
        cell_para(cells[0].cells[0], "1  PUMPHÖLJE   ← IN: hölje", 11, bold=True, center=True)
        cell_para(cells[1].cells[0], "MAGNETVENTIL", 10, bold=True, center=True)
        cell_para(cells[2].cells[0], "~~~~ vatten ~~~~", 10, center=True)
        cell_para(cells[3].cells[0], "2  HAND   ← KROPP", 11, bold=True, center=True)
        cell_para(cells[4].cells[0], "SKROV   3 UT", 11, bold=True, center=True)
    else:
        cell_para(cells[0].cells[0], "PUMPHÖLJE", 11, bold=True, center=True)
        cell_para(cells[1].cells[0], "MAGNETVENTIL", 10, bold=True, center=True)
        cell_para(cells[2].cells[0], "~~~~ vatten ~~~~", 10, center=True)
        cell_para(cells[3].cells[0], "HAND", 11, bold=True, center=True)
        cell_para(cells[4].cells[0], "SKROV", 11, bold=True, center=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_chrome_docx(doc: Document, m: str, mname: str, v: str, vname: str, course: str):
    table = doc.add_table(rows=1, cols=5)
    set_table_borders(table, "0")
    # remove borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)
    c0, c1, c2, c3, c4 = table.rows[0].cells
    shade_cell(c0, INK)
    cell_para(c0, m, 9, bold=True, white=True, center=True)
    cell_para(c1, mname, 11, bold=True)
    shade_cell(c2, INK)
    cell_para(c2, v, 9, bold=True, white=True, center=True)
    cell_para(c3, vname, 11, bold=True)
    cell_para(c4, course, 9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


def add_teacher_banner(doc: Document):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "0")
    c = table.rows[0].cells[0]
    shade_cell(c, INK)
    cell_para(c, "TILL LÄRARE  ·  INTE TILL ELEV  ·  INTE PÅ ELEV-GITHUB PAGES", 9, bold=True, white=True, center=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_meta_fields(doc: Document):
    add_p(doc, "Namn: ________________________________  Datum: __________")
    add_p(doc, "Spår:  ☐ elektriker (land)     ☐ ingenjör (fartyg)     ☐ båda")


def add_megger_table(doc: Document, filled: bool):
    if filled:
        chain_lines = [
            "[1] från",
            "[2] lås / skylt",
            "[3] tvåpol död",
            "[4] megger",
            "[5] protokoll",
        ]
        chain_intro = "Kedja i ordning"
        ref = "☑ skrov / stomme     ☐ PE-skena"
        a = "Beslut: ☑ GO     ☐ STOP\nTillslag: GO — planera tillslag EFTER kåpa på"
        b = "Beslut: ☐ GO     ☑ STOP\nTillslag: STOP — inte spänningssätt, avvikelse"
        ohm = "☐ skall i 5 §     ☑ upplysning"
    else:
        chain_lines = [
            "[ ] megger",
            "[ ] protokoll",
            "[ ] tvåpol död",
            "[ ] från",
            "[ ] lås / skylt",
        ]
        chain_intro = "Kedja — skriv 1–5 i rutorna (rätt ordning)"
        ref = "☐ skrov / stomme     ☐ PE-skena"
        a = "Beslut: ☐ GO     ☐ STOP\nTillslag: ______________________________________"
        b = "Beslut: ☐ GO     ☐ STOP\nTillslag: ______________________________________"
        ohm = "☐ skall i 5 §     ☐ upplysning"

    rows = [
        ("MEGGERKORT", None),
        ("Grupp", "230 V inredning, kabel bytt"),
        ("Tid", "före arbete, avställd grupp"),
        (chain_intro, "\n".join(chain_lines)),
        ("Referens", ref + "\nSamma grupp hela kortet. Referens är skrov/stomme, inte PE-skena."),
        ("Avläsning A\n2,4 MΩ", a),
        ("Avläsning B\n0,4 MΩ", b),
        ("1 MΩ är", ohm),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_borders(table, "12")
    for i, (lab, val) in enumerate(rows):
        left, right = table.rows[i].cells
        if val is None:
            left.merge(right)
            shade_cell(left, INK)
            cell_para(left, lab, 9, bold=True, white=True)
        else:
            cell_para(left, lab, 10, bold=True)
            right.text = ""
            first = True
            for line in val.split("\n"):
                p = right.paragraphs[0] if first else right.add_paragraph()
                first = False
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_inline_runs(p, line, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def setup_docx(teacher: bool, footer: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(16)
        section.right_margin = Mm(16)
        section.top_margin = Mm(14)
        section.bottom_margin = Mm(16)
        footer_p = section.footer.paragraphs[0]
        footer_p.paragraph_format.space_before = Pt(0)
        r = footer_p.add_run(footer)
        set_run_font(r, 8, bold=True)
        if teacher:
            hp = section.header.paragraphs[0]
            r = hp.add_run("LÄRARE — INTE TILL ELEV")
            set_run_font(r, 8, bold=True)
    style = doc.styles["Normal"]
    style.font.name = "IBM Plex Sans"
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK_RGB
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "IBM Plex Sans")
    return doc


def add_picture_if(doc: Document, path: Path, caption: str):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Mm(170))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, 8.5, bold=True)


def write_docx_from_md(md: str, dest: Path, teacher: bool, footer: str,
                       m: str, mname: str, v: str, vname: str,
                       figure: Path | None = None, figure_cap: str = "",
                       post_box: str | None = None):
    doc = setup_docx(teacher, footer)
    if teacher:
        add_teacher_banner(doc)
    add_chrome_docx(doc, m, mname, v, vname, "Elteknik och ellära 45 p")
    blocks = parse_md(md)
    inserted_fig = False
    for b in blocks:
        t = b["type"]
        if t == "h1":
            add_h(doc, b["text"], 18, border_top=False)
            if figure is not None and not inserted_fig:
                add_picture_if(doc, figure, figure_cap)
                inserted_fig = True
        elif t == "h2":
            add_h(doc, b["text"], 13, border_top=True)
        elif t == "h3":
            add_h(doc, b["text"], 11, border_top=False)
        elif t == "p":
            add_p(doc, b["text"])
        elif t == "hr":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "6")
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), INK)
            pBdr.append(top)
            pPr.append(pBdr)
        elif t == "ul":
            for item in b["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.clear()
                add_inline_runs(p, item, 10.5)
        elif t == "ol":
            for item in b["items"]:
                p = doc.add_paragraph(style="List Number")
                p.clear()
                add_inline_runs(p, item, 10.5)
        elif t == "table":
            rows = [b["header"]] + b["rows"]
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            set_table_borders(table, "8")
            for ri, row in enumerate(rows):
                for ci in range(cols):
                    val = row[ci] if ci < len(row) else ""
                    cell_para(table.rows[ri].cells[ci], val, 9.5, bold=(ri == 0))
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
        elif t == "callout":
            add_callout_table(doc, b["title"], b["body"], fara=b["kind"] == "fara")
        elif t == "pre":
            add_p(doc, b["text"], size=9)
    if post_box:
        table = doc.add_table(rows=1, cols=1)
        set_table_borders(table, "8")
        c = table.rows[0].cells[0]
        c.text = ""
        p = c.paragraphs[0]
        add_inline_runs(p, post_box, 11)
        add_p(doc, "Klistra in som det står. Inte studentblad.")
    doc.save(dest)


def write_elev_stotvag_docx(dest: Path):
    doc = setup_docx(False, "Elteknik och ellära 45 p · V1 · elevblad stötväg")
    add_chrome_docx(doc, "M01", "Elsäkerhet", "V1", "Vecka 1", "Elteknik och ellära 45 p")
    add_h(doc, "Stötväg — elevblad", 18)
    add_p(doc, "På papper. Du går inte till luckan. Du megger inte.")
    add_meta_fields(doc)
    add_callout_table(doc, "BILJETT", "Det läcker vid pumpen, känn efter var.")
    add_p(doc, "Pannrum. Pump och magnetventil. Stålskrov.")
    add_callout_table(doc, "OBS", "Peka stötväg innan du tar i något.")
    add_h(doc, "1. STOP eller GO", 13, border_top=True)
    add_p(doc, "Lappen säger: *känn efter var.* Kryssa **en**. Sen ritar du.")
    add_choice_row(doc, ("STOP", "☐ STOP", False), ("GO", "☐ GO", False))
    add_h(doc, "2. Figur — rita stötvägen", 13, border_top=True)
    add_p(doc, "Figuren är blank. Rita stötvägen: **hölje → kropp → skrov**.")
    add_diag(doc, filled=False)
    add_p(doc, "Stötväg (fyll):  ________ →  ________ →  ________")
    add_h(doc, "3. Ring två fel", 13, border_top=True)
    add_p(doc, "Tre lappar. Ring **två** fel.")
    add_three_cards(doc, ["packning", "PE", "utanför IR"], [False, False, False])
    add_h(doc, "4. Två spår — en rad", 13, border_top=True)
    add_p(doc, "Elektriker: leta inte JFB.")
    add_p(doc, "Ingenjör: läcker det vid spänningssatt don är det elarbete.")
    doc.save(dest)


def write_facit_stotvag_docx(dest: Path):
    doc = setup_docx(True, "Elteknik och ellära 45 p · V1 · LÄRARE facit stötväg")
    add_teacher_banner(doc)
    add_chrome_docx(doc, "M01", "Elsäkerhet", "V1", "Vecka 1", "Elteknik och ellära 45 p")
    add_h(doc, "Stötväg — facit", 18)
    add_p(doc, "**G:** STOP, stötväg hölje → kropp → skrov, två av tre fel.")
    add_p(doc, "**IG:** GO, handen först, JFB som skydd.")
    add_h(doc, "1. STOP markerat", 13, border_top=True)
    add_callout_table(doc, "FARA", "Känn inte efter på spänningssatt don.", fara=True)
    add_choice_row(doc, ("STOP", "☑ STOP", True), ("GO", "☐ GO", False))
    add_p(doc, "**Varför (en rad):** våt magnetventil kan bära spänning. Stötväg hölje → kropp → skrov.")
    add_p(doc, "GO = IG.")
    add_h(doc, "2. Figur — stötväg ifylld", 13, border_top=True)
    add_diag(doc, filled=True)
    add_p(doc, "Stötväg: **hölje → kropp → skrov**")
    for t in ("In i höljet.", "Genom kroppen.", "Ut i skrovet."):
        p = doc.add_paragraph(style="List Number")
        p.clear()
        add_inline_runs(p, t, 10.5)
    add_h(doc, "3. Två fel ringade", 13, border_top=True)
    add_p(doc, "**G = två av tre.** Alla tre duger. Visa minst dessa två:")
    add_three_cards(doc, ["packning", "PE", "utanför IR"], [True, True, False])
    add_p(doc, "Tredje (utanför IR) räknas också. IG: inget ringat, eller JFB som “fel”.")
    add_h(doc, "4. Två spår — en rad", 13, border_top=True)
    add_p(doc, "Elektriker = leta inte JFB.")
    add_p(doc, "Ingenjör = läcker det vid spänningssatt don är det elarbete.")
    add_p(doc, "Samma rad på båda spåren. Inte två prov.")
    add_p(doc, "Inte skolessä. Inte 1 MΩ. Inte megger.")
    doc.save(dest)


def write_elev_megger_docx(dest: Path):
    doc = setup_docx(False, "Elteknik och ellära 45 p · V2 · elevblad meggerkort")
    add_chrome_docx(doc, "M02", "Isolering", "V2", "Vecka 2", "Elteknik och ellära 45 p")
    add_h(doc, "Meggerkort — elevblad", 18)
    add_p(doc, "Isolera på papper. Megga inte på spänning.")
    add_meta_fields(doc)
    add_callout_table(
        doc,
        "BILJETT",
        "230 V-grupp i inredningen. Kabel bytt. Isolera på papper. Megga inte på spänning.",
    )
    add_callout_table(doc, "FARA", "Ingen megger på spänningssatt grupp.", fara=True)
    add_h(doc, "Meggerkort (blankt)", 13, border_top=True)
    add_megger_table(doc, filled=False)
    add_h(doc, "Två spår — en rad", 13, border_top=True)
    add_p(doc, "Elektriker: referens är inte PE-skenan hemma.")
    add_p(doc, "Ingenjör: bara för att maskinen inte låter betyder det inte att den är avställd.")
    add_p(doc, "Skriv raden för ditt spår:")
    add_p(doc, "________________________________________________________________")
    add_h(doc, "VG", 13, border_top=True)
    add_p(doc, "Saknas megger på varvet →")
    add_p(doc, "________________________________________________________________")
    add_p(doc, "(tvåpol + …)")
    doc.save(dest)


def write_facit_megger_docx(dest: Path):
    doc = setup_docx(True, "Elteknik och ellära 45 p · V2 · LÄRARE facit meggerkort")
    add_teacher_banner(doc)
    add_chrome_docx(doc, "M02", "Isolering", "V2", "Vecka 2", "Elteknik och ellära 45 p")
    add_h(doc, "Meggerkort — facit", 18)
    add_p(doc, "**G:** kedja i ordning, referens skrov, 2,4 = GO efter kåpa, 0,4 = STOP + avvikelse, 1 MΩ = upplysning.")
    add_p(doc, "**VG:** saknas megger på varvet → tvåpol + gapflagga.")
    add_p(doc, "**IG:** megger på spänning, PE som referens, 1 MΩ som lag, hoppar lås, 0,4 tillslag.")
    add_h(doc, "Arbetsorder", 13, border_top=True)
    add_callout_table(
        doc,
        "BILJETT",
        "230 V-grupp i inredningen. Kabel bytt. Isolera på papper. Megga inte på spänning.",
    )
    add_callout_table(doc, "FARA", "Ingen megger på spänningssatt grupp.", fara=True)
    add_h(doc, "Meggerkort — ifyllt", 13, border_top=True)
    add_p(doc, "Kedja **i ordning**. Referens **skrov/stomme**, inte PE.")
    add_megger_table(doc, filled=True)
    add_p(doc, "**2,4 MΩ** = GO att *planera* tillslag **efter kåpa på**. Inte tillslag med öppen kåpa.")
    add_p(doc, "**0,4 MΩ** = STOP. Ingen tillslag. Skriv avvikelse.")
    add_p(doc, "**1 MΩ** = upplysning, inte lag. Inte “olagligt enligt 5 §”.")
    add_h(doc, "Två spår — en rad", 13, border_top=True)
    add_p(doc, "Elektriker = referens är inte PE-skenan hemma.")
    add_p(doc, "Ingenjör = bara för att maskinen inte låter betyder det inte att den är avställd.")
    add_h(doc, "VG", 13, border_top=True)
    add_p(doc, "Saknas megger på varvet → **tvåpol + gapflagga**. Skriv att meggern saknas. Hitta inte på ett tal. Skippa inte isolation.")
    add_p(doc, "Inte skolessä. Inte 1 MΩ som skall.")
    doc.save(dest)


# ---------------------------------------------------------------------------
# Weeks 3–9 (generic lektion + inlägg; elevblad/facit only if sources exist)
# ---------------------------------------------------------------------------

COURSE = "Elteknik och ellära 45 p"


def doc_lektion(
    md_path: Path | str,
    figure_path: Path | str | None,
    m: str,
    mname: str,
    v: str,
    vname: str,
    footer: str,
    fig_cap: str,
) -> tuple[str, str, bool]:
    md = Path(md_path).read_text(encoding="utf-8")
    body = chrome(m, mname, v, vname, COURSE)
    body += md_body_html(md)
    if figure_path:
        fig = Path(figure_path)
        fig_html = f"""
<div class="fig">
  <img src="{html.escape(fig.resolve().as_uri())}" alt="{html.escape(fig_cap)}"/>
  <div class="cap">{html.escape(fig_cap)}</div>
</div>"""
        if "</h1>" in body:
            body = body.replace("</h1>", "</h1>" + fig_html, 1)
        else:
            body = fig_html + body
    return body, footer, False


def doc_inlagg_from(path: Path | str, m: str, mname: str, v: str, vname: str, footer: str) -> tuple[str, str, bool]:
    md = Path(path).read_text(encoding="utf-8")
    parts = re.split(r"\n---\n", md, maxsplit=1)
    head = parts[0]
    post = parts[1].strip() if len(parts) > 1 else ""
    body = chrome(m, mname, v, vname, COURSE)
    body += md_body_html(head)
    if post:
        body += f'<div class="post">{inline_html(post)}</div>'
    body += '<p style="margin-top:10px"><em>Klistra in som det står. Inte studentblad.</em></p>'
    return body, footer, True


def doc_md(
    src: Path,
    m: str,
    mname: str,
    v: str,
    vname: str,
    footer: str,
    teacher: bool,
) -> tuple[str, str, bool]:
    md = src.read_text(encoding="utf-8")
    body = chrome(m, mname, v, vname, COURSE)
    body += md_body_html(md)
    return body, footer, teacher


def L(src, out, m, mname, figure, figure_cap, footer):
    return {
        "src": ROOT / src,
        "out": out,
        "m": m,
        "mname": mname,
        "figure": ROOT / figure if figure else None,
        "figure_cap": figure_cap,
        "footer": footer,
    }


def P(src, out, m, mname, footer):
    return {
        "src": ROOT / src,
        "out": out,
        "m": m,
        "mname": mname,
        "footer": footer,
    }


WEEK_EXTEND = [
    {
        "week": 3,
        "dir": "vecka-03",
        "v": "V3",
        "vname": "Vecka 3",
        "heading": "Vecka 3 — V3 · M03 Resistiv DC",
        "blurb": "V3 · M03 Resistiv DC. A4. Räkna och mät på papper. Inte händer-DMM som läxa. Inte trefas.",
        "lessons": [
            L(
                "kurs/lektioner/3.1-dc.md",
                "lektion-3.1-dc",
                "M03",
                "Resistiv DC",
                "bok/figur-3-1-dc-matning.png",
                "Figur 3.1 DC-mätning (bok). Serie/parallell. Inte trefas. Inte live tavla.",
                "Elteknik och ellära 45 p · V3 · lektion 3.1 DC",
            ),
        ],
        "posts": [
            P(
                "kurs/lararhandledning/classroom-v3-vecka3.md",
                "larare-inlagg-vecka3",
                "M03",
                "Resistiv DC",
                "Elteknik och ellära 45 p · V3 · LÄRARE inlägg",
            ),
        ],
    },
    {
        "week": 4,
        "dir": "vecka-04",
        "v": "V4",
        "vname": "Vecka 4",
        "heading": "Vecka 4 — V4 · M04 Enfas AC",
        "blurb": "V4 · M04 Enfas AC. A4. rms mot topp. Inte trefas. Inte 440 V. Inte live tavla.",
        "lessons": [
            L(
                "kurs/lektioner/4.1-enfas-ac.md",
                "lektion-4.1-enfas-ac",
                "M04",
                "Enfas AC",
                "bok/figur-4-1-enfas-ac.png",
                "Figur 4.1 Enfas AC (bok). rms mot topp. Inte trefas. Inte 440 V.",
                "Elteknik och ellära 45 p · V4 · lektion 4.1 enfas AC",
            ),
        ],
        "posts": [
            P(
                "kurs/lararhandledning/classroom-v4-vecka4.md",
                "larare-inlagg-vecka4",
                "M04",
                "Enfas AC",
                "Elteknik och ellära 45 p · V4 · LÄRARE inlägg",
            ),
        ],
    },
    {
        "week": 5,
        "dir": "vecka-05",
        "v": "V5",
        "vname": "Vecka 5",
        "heading": "Vecka 5 — V5 · M05 Trefas",
        "blurb": "V5 · M05 Trefas. A4. Linje mot fas. 440 V = look-not-touch. Inte live tavla.",
        "lessons": [
            L(
                "kurs/lektioner/5.1-trefas.md",
                "lektion-5.1-trefas",
                "M05",
                "Trefas",
                "bok/figur-5-1-trefas-spanning.png",
                "Figur 5.1 Trefas (bok). Linje mot fas. 440 V = look-not-touch.",
                "Elteknik och ellära 45 p · V5 · lektion 5.1 trefas",
            ),
        ],
        "posts": [
            P(
                "kurs/lararhandledning/classroom-v5-vecka5.md",
                "larare-inlagg-vecka5",
                "M05",
                "Trefas",
                "Elteknik och ellära 45 p · V5 · LÄRARE inlägg",
            ),
        ],
    },
    {
        "week": 6,
        "dir": "vecka-06",
        "v": "V6",
        "vname": "Vecka 6",
        "heading": "Vecka 6 — V6 · M06 Maskiner",
        "blurb": "V6 · M06 Maskiner. A4. Motor/generator. Inte öppna 440-kåpa. Inte live MSB.",
        "lessons": [
            L(
                "kurs/lektioner/6.1-maskiner.md",
                "lektion-6.1-maskiner",
                "M06",
                "Maskiner",
                "bok/figur-6-1-maskiner.png",
                "Figur 6.1 Maskiner (bok). Motor/generator. Inte öppna 440-kåpa.",
                "Elteknik och ellära 45 p · V6 · lektion 6.1 maskiner",
            ),
        ],
        "posts": [
            P(
                "kurs/lararhandledning/classroom-v6-vecka6.md",
                "larare-inlagg-vecka6",
                "M06",
                "Maskiner",
                "Elteknik och ellära 45 p · V6 · LÄRARE inlägg",
            ),
        ],
    },
    {
        "week": 7,
        "dir": "vecka-07",
        "v": "V7",
        "vname": "Vecka 7",
        "heading": "Vecka 7 — V7 · M07 Eltavla + M08 Verktyg",
        "blurb": "V7 · M07 Eltavla + M08 Verktyg. A4. Look-not-touch på 440. Inte live MSB. Inte megger på spänning.",
        "ticket_m": "M07",
        "ticket_mname": "Eltavla + verktyg",
        "lessons": [
            L(
                "kurs/lektioner/7.1-eltavla.md",
                "lektion-7.1-eltavla",
                "M07",
                "Eltavla",
                "bok/figur-7-1-eltavla.png",
                "Figur 7.1 Eltavla (bok). Huvudtavla vs grupptavla. Inte live MSB.",
                "Elteknik och ellära 45 p · V7 · lektion 7.1 eltavla",
            ),
            L(
                "kurs/lektioner/8.1-verktyg.md",
                "lektion-8.1-verktyg",
                "M08",
                "Verktyg",
                "bok/figur-8-1-verktyg.png",
                "Figur 8.1 Verktyg (bok). DMM, tvåpol, megger. Inte megger på spänning.",
                "Elteknik och ellära 45 p · V7 · lektion 8.1 verktyg",
            ),
        ],
        "posts": [
            P(
                "kurs/lararhandledning/classroom-v7-vecka7.md",
                "larare-inlagg-vecka7",
                "M07",
                "Eltavla",
                "Elteknik och ellära 45 p · V7 · LÄRARE inlägg 7.1",
            ),
            P(
                "kurs/lararhandledning/classroom-v7b-vecka7.md",
                "larare-inlagg-vecka7b",
                "M08",
                "Verktyg",
                "Elteknik och ellära 45 p · V7 · LÄRARE inlägg 8.1",
            ),
        ],
    },
    {
        "week": 8,
        "dir": "vecka-08",
        "v": "V8",
        "vname": "Vecka 8",
        "heading": "Vecka 8 — V8 · M09 Ritningar + M10 Hållkrets",
        "blurb": "V8 · M09 Ritningar + M10 Hållkrets. A4. Schema innan durk. Inte Arduino. Inte live 440.",
        "ticket_m": "M09",
        "ticket_mname": "Ritningar + hållkrets",
        "lessons": [
            L(
                "kurs/lektioner/9.1-ritningar.md",
                "lektion-9.1-ritningar",
                "M09",
                "Ritningar",
                "bok/figur-9-1-ritningar.png",
                "Figur 9.1 Ritningar (bok). Enlinje mot kretsschema. Inte gissa live.",
                "Elteknik och ellära 45 p · V8 · lektion 9.1 ritningar",
            ),
            L(
                "kurs/lektioner/10.1-hallkrets.md",
                "lektion-10.1-hallkrets",
                "M10",
                "Hållkrets",
                "bok/figur-10-1-hallkrets.png",
                "Figur 10.1 Hållkrets (bok). Start–håll–stopp. Inte Arduino. Inte live 440.",
                "Elteknik och ellära 45 p · V8 · lektion 10.1 hållkrets",
            ),
        ],
        "posts": [
            P(
                "kurs/lararhandledning/classroom-v8-vecka8.md",
                "larare-inlagg-vecka8",
                "M09",
                "Ritningar",
                "Elteknik och ellära 45 p · V8 · LÄRARE inlägg 9.1",
            ),
            P(
                "kurs/lararhandledning/classroom-v8b-vecka8.md",
                "larare-inlagg-vecka8b",
                "M10",
                "Hållkrets",
                "Elteknik och ellära 45 p · V8 · LÄRARE inlägg 10.1",
            ),
        ],
    },
    {
        "week": 9,
        "dir": "vecka-09",
        "v": "V9",
        "vname": "Vecka 9",
        "heading": "Vecka 9 — V9 · M11 Elarbete + M12 Felsökning",
        "blurb": "V9 · M11 Elarbete + M12 Felsökning. A4. Inte TS-blankett. Inte PE-jakt. Inte live 440. Facit är läraronly (filnamn *facit*). Inte quiz.",
        "ticket_m": "M11",
        "ticket_mname": "Elarbete + felsökning",
        "lessons": [
            L(
                "kurs/lektioner/11.1-arbete-ip-intyg.md",
                "lektion-11.1-arbete-ip-intyg",
                "M11",
                "Elarbete",
                "bok/figur-11-1-arbete-ip-intyg.png",
                "Figur 11.1 Elarbete, IP, intyg (bok). Inte TS-blankett. Inte 440-lucka.",
                "Elteknik och ellära 45 p · V9 · lektion 11.1 elarbete",
            ),
            L(
                "kurs/lektioner/12.1-felsokning.md",
                "lektion-12.1-felsokning",
                "M12",
                "Felsökning",
                "bok/figur-12-1-felsokning.png",
                "Figur 12.1 Felsökning (bok). Kedja på papper. Inte PE-jakt. Inte live 440.",
                "Elteknik och ellära 45 p · V9 · lektion 12.1 felsökning",
            ),
        ],
        "posts": [
            P(
                "kurs/lararhandledning/classroom-v9-vecka9.md",
                "larare-inlagg-vecka9",
                "M11",
                "Elarbete",
                "Elteknik och ellära 45 p · V9 · LÄRARE inlägg 11.1",
            ),
            P(
                "kurs/lararhandledning/classroom-v9b-vecka9.md",
                "larare-inlagg-vecka9b",
                "M12",
                "Felsökning",
                "Elteknik och ellära 45 p · V9 · LÄRARE inlägg 12.1",
            ),
            P(
                "kurs/lararhandledning/classroom-v-prov-vecka9.md",
                "larare-inlagg-prov-vecka9",
                "PROV",
                "Skriftligt prov",
                "Elteknik och ellära 45 p · V9 · LÄRARE inlägg skriftligt prov",
            ),
        ],
    },
]


def elevblad_sources(week: int) -> list[tuple[Path, Path, bool, str]]:
    """Return (src, dest_stem, teacher, slug) for existing elevblad this week.

    Never invent a ticket. Facit never uses elev-* filenames.
    """
    folder = ROOT / "kurs" / "elevblad"
    found = []
    for elev in sorted(folder.glob(f"v{week}-*-elev.md")):
        slug = elev.name[len(f"v{week}-") : -len("-elev.md")]
        found.append((elev, f"elev-{slug}", False, slug))
    for facit in sorted(folder.glob(f"v{week}-*-facit.md")):
        slug = facit.name[len(f"v{week}-") : -len("-facit.md")]
        found.append((facit, f"larare-{slug}-facit", True, slug))
    return found


def write_week_readme(spec: dict, elev_names: list[str], teacher_names: list[str]) -> None:
    dest = OUT / spec["dir"] / "README.md"
    elev_line = ", ".join(f"`{n}.pdf` + `.docx`" for n in elev_names) if elev_names else "ingen (inget elevblad i `kurs/elevblad/`)"
    teacher_line = ", ".join(f"`{n}.pdf` + `.docx`" for n in teacher_names) if teacher_names else "—"
    dest.write_text(
        f"# {spec['vname']} — Classroom\n\n"
        f"**Elev:** {elev_line}\n\n"
        f"**Lärare (inte till elev):** {teacher_line}\n\n"
        f"{spec['blurb']}\n",
        encoding="utf-8",
    )


def build_week_extend(spec: dict) -> None:
    week_dir = OUT / spec["dir"]
    week_dir.mkdir(parents=True, exist_ok=True)
    v, vname = spec["v"], spec["vname"]
    elev_stems: list[str] = []
    teacher_stems: list[str] = []

    for lesson in spec["lessons"]:
        src = lesson["src"]
        if not src.exists():
            raise FileNotFoundError(src)
        stem = lesson["out"]
        body, footer, teacher = doc_lektion(
            src,
            lesson["figure"],
            lesson["m"],
            lesson["mname"],
            v,
            vname,
            lesson["footer"],
            lesson["figure_cap"],
        )
        pdf = week_dir / f"{stem}.pdf"
        docx = week_dir / f"{stem}.docx"
        print(f"PDF  {pdf.relative_to(ROOT)}")
        write_pdf(body, footer, teacher, pdf)
        write_docx_from_md(
            src.read_text(encoding="utf-8"),
            docx,
            teacher=False,
            footer=footer,
            m=lesson["m"],
            mname=lesson["mname"],
            v=v,
            vname=vname,
            figure=lesson["figure"],
            figure_cap=lesson["figure_cap"],
        )
        print(f"DOCX {docx.relative_to(ROOT)}")
        elev_stems.append(stem)

    for post in spec["posts"]:
        src = post["src"]
        if not src.exists():
            print(f"SKIP missing post {src.relative_to(ROOT)}")
            continue
        stem = post["out"]
        body, footer, teacher = doc_inlagg_from(src, post["m"], post["mname"], v, vname, post["footer"])
        pdf = week_dir / f"{stem}.pdf"
        docx = week_dir / f"{stem}.docx"
        print(f"PDF  {pdf.relative_to(ROOT)}")
        write_pdf(body, footer, teacher, pdf)
        raw = src.read_text(encoding="utf-8")
        parts = re.split(r"\n---\n", raw, maxsplit=1)
        head = parts[0]
        box = parts[1].strip() if len(parts) > 1 else ""
        write_docx_from_md(
            head,
            docx,
            teacher=True,
            footer=footer,
            m=post["m"],
            mname=post["mname"],
            v=v,
            vname=vname,
            post_box=box,
        )
        print(f"DOCX {docx.relative_to(ROOT)}")
        teacher_stems.append(stem)

    if spec["week"] == 9:
        exam_jobs = [
            {
                "src": ROOT / "kurs/prov/skriftligt-prov-elev.md",
                "out": "elev-skriftligt-prov",
                "m": "PROV",
                "mname": "Skriftligt prov",
                "footer": "Elteknik och ellära 45 p · V9 · Elevblad · skriftligt prov",
                "teacher": False,
            },
            {
                "src": ROOT / "kurs/prov/skriftligt-prov-facit.md",
                "out": "larare-skriftligt-prov-facit",
                "m": "PROV",
                "mname": "Skriftligt prov",
                "footer": "Elteknik och ellära 45 p · V9 · LÄRARE facit skriftligt prov",
                "teacher": True,
            },
        ]
        for job in exam_jobs:
            src = job["src"]
            if not src.exists():
                raise FileNotFoundError(src)
            if job["teacher"] and "facit" not in job["out"]:
                raise RuntimeError(f"exam facit filename must contain facit: {job['out']}")
            if (not job["teacher"]) and ("facit" in job["out"] or job["out"].startswith("larare-")):
                raise RuntimeError(f"student exam must not be teacher-named: {job['out']}")
            body, footer, teacher_flag = doc_md(
                src,
                job["m"],
                job["mname"],
                v,
                vname,
                job["footer"],
                job["teacher"],
            )
            pdf = week_dir / f"{job['out']}.pdf"
            docx = week_dir / f"{job['out']}.docx"
            print(f"PDF  {pdf.relative_to(ROOT)}")
            write_pdf(body, footer, teacher_flag, pdf)
            write_docx_from_md(
                src.read_text(encoding="utf-8"),
                docx,
                teacher=teacher_flag,
                footer=footer,
                m=job["m"],
                mname=job["mname"],
                v=v,
                vname=vname,
            )
            print(f"DOCX {docx.relative_to(ROOT)}")
            (teacher_stems if job["teacher"] else elev_stems).append(job["out"])

    fallback_m = spec.get("ticket_m") or (spec["lessons"][0]["m"] if spec["lessons"] else "M")
    fallback_name = spec.get("ticket_mname") or (spec["lessons"][0]["mname"] if spec["lessons"] else "")
    for src, stem, teacher, slug in elevblad_sources(spec["week"]):
        if stem.startswith("elev-") and "facit" in stem:
            raise RuntimeError(f"facit must not use elev-* name: {stem}")
        footer = (
            f"{COURSE} · {v} · LÄRARE facit {slug}"
            if teacher
            else f"{COURSE} · {v} · elevblad {slug}"
        )
        body, footer, teacher_flag = doc_md(
            src, fallback_m, fallback_name, v, vname, footer, teacher
        )
        pdf = week_dir / f"{stem}.pdf"
        docx = week_dir / f"{stem}.docx"
        print(f"PDF  {pdf.relative_to(ROOT)}")
        write_pdf(body, footer, teacher_flag, pdf)
        write_docx_from_md(
            src.read_text(encoding="utf-8"),
            docx,
            teacher=teacher_flag,
            footer=footer,
            m=fallback_m,
            mname=fallback_name,
            v=v,
            vname=vname,
        )
        print(f"DOCX {docx.relative_to(ROOT)}")
        (teacher_stems if teacher else elev_stems).append(stem)

    elev_stems.sort(key=lambda s: 0 if s.startswith("elev-") else 1)
    teacher_stems.sort(key=lambda s: 0 if "facit" in s else 1)
    write_week_readme(spec, elev_stems, teacher_stems)


def build_all():
    (OUT / "vecka-01").mkdir(parents=True, exist_ok=True)
    (OUT / "vecka-02").mkdir(parents=True, exist_ok=True)

    pdf_specs = [
        (doc_elev_stotvag, OUT / "vecka-01/elev-stotvag.pdf"),
        (doc_facit_stotvag, OUT / "vecka-01/larare-stotvag-facit.pdf"),
        (doc_lektion_11, OUT / "vecka-01/lektion-1.1-stotar.pdf"),
        (lambda: doc_inlagg(1), OUT / "vecka-01/larare-inlagg-vecka1.pdf"),
        (doc_elev_megger, OUT / "vecka-02/elev-meggerkort.pdf"),
        (doc_facit_megger, OUT / "vecka-02/larare-meggerkort-facit.pdf"),
        (doc_lektion_21, OUT / "vecka-02/lektion-2.1-isolering.pdf"),
        (lambda: doc_inlagg(2), OUT / "vecka-02/larare-inlagg-vecka2.pdf"),
    ]
    for fn, dest in pdf_specs:
        body, footer, teacher = fn()
        print(f"PDF  {dest.relative_to(ROOT)}")
        write_pdf(body, footer, teacher, dest)

    write_elev_stotvag_docx(OUT / "vecka-01/elev-stotvag.docx")
    write_facit_stotvag_docx(OUT / "vecka-01/larare-stotvag-facit.docx")
    write_elev_megger_docx(OUT / "vecka-02/elev-meggerkort.docx")
    write_facit_megger_docx(OUT / "vecka-02/larare-meggerkort-facit.docx")

    md11 = (ROOT / "kurs/lektioner/1.1-stotar.md").read_text(encoding="utf-8")
    write_docx_from_md(
        md11,
        OUT / "vecka-01/lektion-1.1-stotar.docx",
        teacher=False,
        footer="Elteknik och ellära 45 p · V1 · lektion 1.1 stötväg",
        m="M01", mname="Elsäkerhet", v="V1", vname="Vecka 1",
        figure=ROOT / "bok/figur-1-1-stotvag-ventil.png",
        figure_cap="Figur 1.1 Stötväg — magnetventil (bok). Inte megger. Inte 1 MΩ som skall.",
    )
    md21 = (ROOT / "kurs/lektioner/2.1-isolering.md").read_text(encoding="utf-8")
    write_docx_from_md(
        md21,
        OUT / "vecka-02/lektion-2.1-isolering.docx",
        teacher=False,
        footer="Elteknik och ellära 45 p · V2 · lektion 2.1 isolering",
        m="M02", mname="Isolering", v="V2", vname="Vecka 2",
        figure=ROOT / "bok/figur-2-1-isolering-kedja.png",
        figure_cap="Figur 2.1 Isolering före arbete (bok). Kedja från–lås–prova–megger. 1 MΩ = upplysning.",
    )

    for week, src, dest, m, name, v, footer in [
        (1, ROOT / "kurs/lararhandledning/classroom-v1-vecka1.md",
         OUT / "vecka-01/larare-inlagg-vecka1.docx",
         "M01", "Elsäkerhet", "V1",
         "Elteknik och ellära 45 p · V1 · LÄRARE inlägg"),
        (2, ROOT / "kurs/lararhandledning/classroom-v2-vecka2.md",
         OUT / "vecka-02/larare-inlagg-vecka2.docx",
         "M02", "Isolering", "V2",
         "Elteknik och ellära 45 p · V2 · LÄRARE inlägg"),
    ]:
        raw = src.read_text(encoding="utf-8")
        parts = re.split(r"\n---\n", raw, maxsplit=1)
        head = parts[0]
        post = parts[1].strip() if len(parts) > 1 else ""
        write_docx_from_md(
            head,
            dest,
            teacher=True,
            footer=footer,
            m=m, mname=name, v=v, vname=f"Vecka {week}",
            post_box=post,
        )
        print(f"DOCX {dest.relative_to(ROOT)}")

    print("DOCX vecka-01/elev-stotvag.docx")
    print("DOCX vecka-01/larare-stotvag-facit.docx")
    print("DOCX vecka-01/lektion-1.1-stotar.docx")
    print("DOCX vecka-02/elev-meggerkort.docx")
    print("DOCX vecka-02/larare-meggerkort-facit.docx")
    print("DOCX vecka-02/lektion-2.1-isolering.docx")

    for spec in WEEK_EXTEND:
        build_week_extend(spec)


if __name__ == "__main__":
    build_all()
